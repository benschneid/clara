"""
parser.py — Core .docx → structured chunks logic.

Strategy
--------
Contracts are hierarchically organised by section headings. An AI agent
reviewing a contract does NOT want the entire text in one shot; it wants
to navigate to the relevant clause and read only that. We therefore:

  1. Walk every block in the document (paragraphs + tables, in order).
  2. Detect headings via Word styles *and* via a numbered-section regex
     fallback (handles legacy contracts converted from plain text where
     every paragraph has the same style).
  3. Accumulate non-heading blocks under the current heading, grouping
     consecutive list items and preserving tables as Markdown.
  4. Emit a flat list of Chunk objects, each carrying:
       - a stable ID  (slugified heading + index)
       - a breadcrumb path  (["Section 9", "Section 9.1"])
       - the content as plain text / Markdown
       - a token count estimate (cl100k, approximated without tiktoken)
       - the element type  (paragraph | list | table | mixed)

The response to POST /documents contains:
  - doc_id  (UUID)
  - title   (detected from Title style or first non-blank line)
  - manifest  (list of {id, path, heading, token_count, element_types})
              — NO content, so it is always small regardless of doc size

Individual chunks are fetched via GET /documents/{doc_id}/chunks/{chunk_id}.
This lets an agent load the map cheaply and pull only what it needs.
"""

from __future__ import annotations

import re
import unicodedata
import uuid
from dataclasses import dataclass, field, asdict
from typing import Optional
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------

HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3", "heading 4", "heading 5",
    "heading 6", "title",
}

# Matches: "1.  Foo", "1.1 Foo", "12.3.4 Foo", "EXHIBIT A", "SCHEDULE 1"
SECTION_RE = re.compile(
    r"^(\d+(?:\.\d+)*\.?\s{1,6}\S|EXHIBIT\s+\w|SCHEDULE\s+\w|ANNEX\s+\w)",
    re.IGNORECASE,
)


def _style_name(para) -> str:
    return para.style.name.lower()


def _is_heading(para) -> bool:
    """Return True if this paragraph should open a new chunk."""
    name = _style_name(para)
    if any(name.startswith(h) for h in HEADING_STYLES):
        return True
    text = para.text.strip()
    if not text:
        return False
    # Numbered-section heuristic for plain-style documents
    if SECTION_RE.match(text) and len(text) < 120:
        return True
    return False


def _heading_level(para) -> int:
    """Approximate heading depth for breadcrumb building."""
    name = _style_name(para)
    for level in range(1, 7):
        if f"heading {level}" in name:
            return level
    if "title" in name:
        return 0
    text = para.text.strip()
    # Count dots in leading numeric segment: "1." → 1, "1.1" → 2, "1.1.1" → 3
    m = re.match(r"^(\d+(?:\.\d+)*)", text)
    if m:
        return m.group(1).count(".") + 1
    return 1


# ---------------------------------------------------------------------------
# Table → Markdown
# ---------------------------------------------------------------------------

def _table_to_markdown(table) -> str:
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


# ---------------------------------------------------------------------------
# Slug / ID helpers
# ---------------------------------------------------------------------------

def _slugify(text: str) -> str:
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode()
    text = re.sub(r"[^\w\s-]", "", text).strip().lower()
    text = re.sub(r"[\s_-]+", "-", text)
    return text[:60]


def _chunk_id(heading: str, index: int) -> str:
    slug = _slugify(heading) or "chunk"
    return f"{slug}-{index}"


# ---------------------------------------------------------------------------
# Token counting (approximation — avoids tiktoken dependency)
# ---------------------------------------------------------------------------

def _approx_tokens(text: str) -> int:
    """~4 chars per token is a good rule of thumb for English legal prose."""
    return max(1, len(text) // 4)


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class Chunk:
    id: str
    path: list[str]          # breadcrumb, e.g. ["9. Confidentiality", "9.2 InterTrust Information"]
    heading: str
    content: str
    element_types: list[str] # e.g. ["paragraph", "list", "table"]
    token_count: int

    def manifest_entry(self) -> dict:
        """Lightweight view — no content — for the manifest."""
        return {
            "id": self.id,
            "path": self.path,
            "heading": self.heading,
            "token_count": self.token_count,
            "element_types": self.element_types,
        }

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class ParsedDocument:
    doc_id: str
    title: str
    chunks: list[Chunk]
    total_tokens: int
    chunk_count: int

    def manifest(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "title": self.title,
            "chunk_count": self.chunk_count,
            "total_tokens": self.total_tokens,
            "chunks": [c.manifest_entry() for c in self.chunks],
        }

    def get_chunk(self, chunk_id: str) -> Optional[Chunk]:
        for c in self.chunks:
            if c.id == chunk_id:
                return c
        return None


# ---------------------------------------------------------------------------
# Main parse function
# ---------------------------------------------------------------------------

def parse_docx(path: str, doc_id: Optional[str] = None) -> ParsedDocument:
    """
    Parse a .docx file and return a ParsedDocument with structured chunks.

    Parameters
    ----------
    path    : filesystem path to the .docx file
    doc_id  : optional UUID; one is generated if not supplied
    """
    doc_id = doc_id or str(uuid.uuid4())
    document = Document(path)

    # ---- build a flat list of "blocks" (paragraph or table) in order ----
    # We iterate document.element children to preserve table/paragraph order.
    blocks = []
    body = document.element.body
    para_map = {p._element: p for p in document.paragraphs}
    table_map = {t._element: t for t in document.tables}

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p" and child in para_map:
            blocks.append(("para", para_map[child]))
        elif tag == "tbl" and child in table_map:
            blocks.append(("table", table_map[child]))

    # ---- detect document title ----
    title = "Untitled Document"
    for kind, obj in blocks:
        if kind == "para":
            name = _style_name(obj)
            if "title" in name and obj.text.strip():
                title = obj.text.strip()
                break
            if obj.text.strip() and len(obj.text.strip()) < 120:
                title = obj.text.strip()
                break

    # ---- group blocks into chunks ----
    chunks: list[Chunk] = []
    current_heading = title
    current_path: list[str] = [title]
    current_parts: list[str] = []
    current_types: list[str] = []
    chunk_index = 0

    # Breadcrumb stack: maps heading level → heading text
    path_stack: dict[int, str] = {0: title}

    def flush(heading, path, parts, types, index):
        content = "\n\n".join(p for p in parts if p.strip())
        if not content.strip():
            return None
        token_count = _approx_tokens(content)
        chunk = Chunk(
            id=_chunk_id(heading, index),
            path=list(path),
            heading=heading,
            content=content,
            element_types=sorted(set(types)),
            token_count=token_count,
        )
        return chunk

    for kind, obj in blocks:
        if kind == "para":
            para = obj
            text = para.text.strip()

            if _is_heading(para):
                # flush current accumulator
                result = flush(current_heading, current_path, current_parts, current_types, chunk_index)
                if result:
                    chunks.append(result)
                    chunk_index += 1

                # update breadcrumb
                level = _heading_level(para)
                # trim path stack to this level
                path_stack = {k: v for k, v in path_stack.items() if k < level}
                path_stack[level] = text
                current_path = [path_stack[k] for k in sorted(path_stack)]
                current_heading = text
                current_parts = []
                current_types = []

            else:
                if not text:
                    continue
                # detect list items
                name = _style_name(para)
                if "list" in name or _is_list_item(para):
                    current_parts.append(f"- {text}")
                    _add_type(current_types, "list")
                else:
                    current_parts.append(text)
                    _add_type(current_types, "paragraph")

        elif kind == "table":
            md = _table_to_markdown(obj)
            current_parts.append(md)
            _add_type(current_types, "table")

    # flush final chunk
    result = flush(current_heading, current_path, current_parts, current_types, chunk_index)
    if result:
        chunks.append(result)

    total_tokens = sum(c.token_count for c in chunks)

    return ParsedDocument(
        doc_id=doc_id,
        title=title,
        chunks=chunks,
        total_tokens=total_tokens,
        chunk_count=len(chunks),
    )


def _is_list_item(para) -> bool:
    """Check for list formatting via numPr XML element."""
    return para._element.find(f".//{qn('w:numPr')}") is not None


def _add_type(types: list, t: str):
    if t not in types:
        types.append(t)
